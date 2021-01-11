from docx import Document
from docx.shared import Inches 
import pyttsx3
document = Document()

def speak(text):
    pyttsx3.speak(text)

# profile picture 
document.add_picture(
    'me.pic.JPG', 
    width=Inches(1.0)
)

# name, phone number and email details
name = input('what is your full name? ')
speak('hello' + name + 'hope you are doing alright')
phone_number = input('what is your phone number? ')
email = input('what is your email? ')

document.add_paragraph(
    name + ' / ' + phone_number + ' / ' + email)

# about me
document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# Academic Backround
document.add_heading('Academic Backround')
p = document.add_paragraph()

curriculum = input('enter curriculum ')
instatution = input('enter istatution')
from_date = input('Frome date ')
to_date = input('To DAte ')

p.add_run(curriculum + ' ' + 'curruculum' + ' ' + '\n').bold = True
p.add_run(instatution + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic


#Extracurricular 
document.add_heading('Extracurricular')
p = document.add_paragraph()

extracurricular = input('enter extracurricular activity')
describtion = input('describe your experience')
p.add_run(extracurricular + ':' + '\n' + describtion + '\n').italic

while True:
    more_extracurricular = input('Enter additional extracurricular? Yes or no ')
    if more_extracurricular.lower() == 'yes':
        extracurricular = input('enter extracurricular activity')
        describtion = input('describe your experience')
        p.add_run(extracurricular + ':' + '\n' + describtion + '\n').italic
    else:
        break




# Skills
document.add_heading('Skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skills = input('Do you have additional skills? Yes or No ')
    if more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break 






    





# footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV program generated using code course project '

document.save('cv.docx')
